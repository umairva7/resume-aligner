import re
from pathlib import Path
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.config import settings
from app.utils.file_helpers import generate_unique_filename

class DOCXGeneratorService:
    """
    OOP Service for compiling Markdown-formatted resumes into ATS-compliant Microsoft Word (.docx) files.
    """

    def create_ats_docx(self, markdown_text: str, filename_prefix: str = "tailored_resume") -> Path:
        output_filename = f"{generate_unique_filename(filename_prefix)}.docx"
        output_path = settings.TAILORED_RESUME_DIR / output_filename
        settings.TAILORED_RESUME_DIR.mkdir(parents=True, exist_ok=True)

        template_path = getattr(settings, 'DOCX_TEMPLATE_PATH', None)
        
        has_template = False
        if template_path and template_path.exists():
            try:
                doc = docx.Document(str(template_path))
                # Clear existing content from the template (retain styles, headers, footers)
                for paragraph in doc.paragraphs:
                    p = paragraph._element
                    p.getparent().remove(p)
                    p._p = p._element = None
                for table in doc.tables:
                    t = table._element
                    t.getparent().remove(t)
                    t._tbl = t._element = None
                has_template = True
            except Exception:
                doc = docx.Document()
        else:
            doc = docx.Document()

        if not has_template:
            # Set 0.5 inch margins for default blank doc
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

        lines = markdown_text.splitlines()

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            if line_str.startswith("# "):
                text = line_str.lstrip("#").strip()
                p = self._add_styled_paragraph(doc, 'Heading 1')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                self._add_formatted_runs(p, text, is_heading=True, level=1)
            elif line_str.startswith("## "):
                text = line_str.lstrip("#").strip()
                p = self._add_styled_paragraph(doc, 'Heading 2')
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                self._add_formatted_runs(p, text, is_heading=True, level=2)
            elif line_str.startswith("### "):
                text = line_str.lstrip("#").strip()
                p = self._add_styled_paragraph(doc, 'Heading 3')
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                self._add_formatted_runs(p, text, is_heading=True, level=3)
            elif line_str.startswith("- ") or line_str.startswith("* "):
                text = line_str[2:].strip()
                p = self._add_styled_paragraph(doc, 'List Bullet')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                self._add_formatted_runs(p, text)
            else:
                p = self._add_styled_paragraph(doc, 'Normal')
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                
                # Center align contact info line
                if "|" in line_str and "@" in line_str:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                self._add_formatted_runs(p, line_str)

        doc.save(str(output_path))
        return output_path

    def _add_styled_paragraph(self, doc, style_name):
        try:
            return doc.add_paragraph(style=style_name)
        except KeyError:
            return doc.add_paragraph()

    def _add_formatted_runs(self, paragraph, text: str, is_heading: bool = False, level: int = 0):
        """Parse basic **bold** tags and [text](url) links and add formatted runs to docx paragraph."""
        from docx.oxml.shared import OxmlElement, qn
        import docx.opc.constants
        
        # Split by markdown links first
        link_tokens = re.split(r'(\[.*?\]\(.*?\))', text)
        for l_token in link_tokens:
            if not l_token:
                continue
                
            if l_token.startswith("[") and "](" in l_token and l_token.endswith(")"):
                # It's a link
                link_text = l_token[1:l_token.index("]")]
                link_url = l_token[l_token.index("(")+1:-1]
                
                # Add hyperlink to paragraph
                part = paragraph.part
                r_id = part.relate_to(link_url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), r_id)
                new_run = OxmlElement('w:r')
                
                rPr = OxmlElement('w:rPr')
                c = OxmlElement('w:color')
                c.set(qn('w:val'), '0563C1') # Standard link blue
                rPr.append(c)
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                
                # Apply font properties
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Calibri')
                rFonts.set(qn('w:hAnsi'), 'Calibri')
                rPr.append(rFonts)
                
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '20') # 10pt = 20 half-points
                rPr.append(sz)
                
                new_run.append(rPr)
                
                t = OxmlElement('w:t')
                t.text = link_text
                new_run.append(t)
                
                hyperlink.append(new_run)
                paragraph._p.append(hyperlink)
                continue

            # Parse bold within non-link text
            tokens = re.split(r'(\*\*.*?\*\*)', l_token)
            for token in tokens:
                if not token:
                    continue
                
                is_bold = False
                run_text = token
                if token.startswith("**") and token.endswith("**"):
                    is_bold = True
                    run_text = token[2:-2]
                    
                run = paragraph.add_run(run_text)
                
                if is_bold:
                    run.bold = True
                    
                # Apply hardcoded professional ATS styles
                run.font.name = 'Calibri'
                if is_heading:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    if level == 1:
                        run.font.size = Pt(18)
                    elif level == 2:
                        run.font.size = Pt(13)
                    elif level == 3:
                        run.font.size = Pt(11)
                else:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
